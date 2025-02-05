import PyPDF2
from docx import Document
from difflib import SequenceMatcher
import hazm
import re
from pathlib import Path

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def clean_text(text):
    """Clean and normalize Farsi text."""
    normalizer = hazm.Normalizer()
    text = normalizer.normalize(text)
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def find_common_and_unique_content(text1, text2, similarity_threshold=0.8):
    """Find common and unique content between two Farsi texts."""
    # Clean and normalize texts
    text1 = clean_text(text1)
    text2 = clean_text(text2)
    
    # Use hazm's sentence tokenizer for Farsi
    sent_tokenizer = hazm.SentenceTokenizer()
    sentences1 = sent_tokenizer.tokenize(text1)
    sentences2 = sent_tokenizer.tokenize(text2)
    
    # Find common and unique sentences
    common_sentences = []
    unique_sentences1 = []
    unique_sentences2 = []
    
    for sent1 in sentences1:
        found_match = False
        for sent2 in sentences2:
            # Using sequence matcher to find similar sentences
            similarity = SequenceMatcher(None, sent1, sent2).ratio()
            if similarity > similarity_threshold:  # Threshold for considering sentences similar
                common_sentences.append(sent1)
                found_match = True
                break
        if not found_match:
            unique_sentences1.append(sent1)
    
    # Find unique sentences in text2
    for sent2 in sentences2:
        is_unique = True
        for sent1 in sentences1:
            similarity = SequenceMatcher(None, sent1, sent2).ratio()
            if similarity > similarity_threshold:
                is_unique = False
                break
        if is_unique:
            unique_sentences2.append(sent2)
    
    return common_sentences, unique_sentences1, unique_sentences2

def calculate_statistics(text1, text2, common_sentences):
    """Calculate comparison statistics for Farsi texts."""
    # Word statistics using hazm
    word_tokenizer = hazm.WordTokenizer()
    words1 = word_tokenizer.tokenize(clean_text(text1))
    words2 = word_tokenizer.tokenize(clean_text(text2))
    common_words = set(words1) & set(words2)
    
    # Sentence statistics
    sent_tokenizer = hazm.SentenceTokenizer()
    sentences1 = sent_tokenizer.tokenize(clean_text(text1))
    sentences2 = sent_tokenizer.tokenize(clean_text(text2))
    
    stats = {
        'total_sentences_doc1': len(sentences1),
        'total_sentences_doc2': len(sentences2),
        'common_sentences': len(common_sentences),
        'total_words_doc1': len(words1),
        'total_words_doc2': len(words2),
        'common_words': len(common_words),
        'commonality_percentage': (len(common_sentences) / 
            ((len(sentences1) + len(sentences2)) / 2)) * 100
    }
    
    return stats

def save_to_word(content, output_path):
    """Save content to a Word document with right-to-left text direction."""
    doc = Document()
    
    # Set RTL paragraph direction for Farsi text
    if isinstance(content, list):
        for item in content:
            paragraph = doc.add_paragraph(item)
            paragraph.paragraph_format.right_to_left = True
    else:
        paragraph = doc.add_paragraph(content)
        paragraph.paragraph_format.right_to_left = True
    
    doc.save(output_path)

def compare_pdfs(pdf1_path, pdf2_path, output_dir):
    """Main function to compare Farsi PDFs and generate output files."""
    # Ensure output directory exists
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Extract text from PDFs
    text1 = extract_text_from_pdf(pdf1_path)
    text2 = extract_text_from_pdf(pdf2_path)
    
    # Find common and unique content
    common_sentences, unique1, unique2 = find_common_and_unique_content(text1, text2)
    
    # Calculate statistics
    stats = calculate_statistics(text1, text2, common_sentences)
    
    # Save results to Word documents
    save_to_word(common_sentences, output_dir / 'common_content.docx')
    save_to_word(unique1, output_dir / 'unique_doc1.docx')
    save_to_word(unique2, output_dir / 'unique_doc2.docx')
    
    # Create statistics report
    stats_report = f"""آمار مقایسه:
    تعداد جملات سند ۱: {stats['total_sentences_doc1']}
    تعداد جملات سند ۲: {stats['total_sentences_doc2']}
    تعداد جملات مشترک: {stats['common_sentences']}
    تعداد کلمات سند ۱: {stats['total_words_doc1']}
    تعداد کلمات سند ۲: {stats['total_words_doc2']}
    تعداد کلمات مشترک: {stats['common_words']}
    درصد اشتراک کلی: {stats['commonality_percentage']:.2f}%
    """
    
    save_to_word(stats_report, output_dir / 'statistics.docx')
    
    return stats

if __name__ == "__main__":
    pdf1_path = r"path\to\pdf1"
    pdf2_path = r"path_to_pdf2"
    output_dir = r"path\to\output"
    stats = compare_pdfs(pdf1_path, pdf2_path, output_dir)
